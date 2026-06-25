"""docgen_word_standard_generate — 基于填充模板生成 Word 文档

将填充后的模板数据转为 Word 文档大纲，调用 wordmaster 的 docx 生成能力。
"""

from __future__ import annotations

import json
import logging
from pathlib import Path
from typing import Any

from ark_agentic.core.tools.base import AgentTool, ToolParameter
from ark_agentic.core.types import AgentToolResult, ToolResultType

from egis_agent_plugins.core.internal.a2ui import (
    FrontendDigest, MinimalView, DetailedView, ViewSection,
    ToolDisplayType, apply_dual_layer,
)
from egis_agent_plugins.core.tools.docgen._project_model import (
    ProjectStore, ArtifactMeta, append_event,
)
from egis_agent_plugins.core.tools.docgen._project_model.artifacts import (
    read_artifact, write_artifact,
)

logger = logging.getLogger(__name__)


class DocgenWordStandardGenerateTool(AgentTool):
    """基于填充模板生成 Word 文档

    读取填充后的模板，将其转为 docx 大纲格式，
    委托 wordmaster 的 docx_generate_from_outline 生成 Word 文档。
    """

    name = "docgen_word_standard_generate"
    description = (
        "基于填充后的投标模板生成 Word 文档初稿。\n"
        "将模板数据转为 docx 大纲格式并生成 .docx 文件。\n"
        "生成后可以使用 Word 编辑器打开查看和编辑。"
    )
    parameters = [
        ToolParameter(
            name="project_path",
            type="string",
            description="项目根目录路径",
            required=True,
        ),
        ToolParameter(
            name="filled_template_artifact_key",
            type="string",
            description="填充模板 artifact 键名。默认: 'filled_template'",
            required=False,
        ),
        ToolParameter(
            name="output_artifact_key",
            type="string",
            description="输出 Word 的 artifact 键名。默认: 'draft_word'",
            required=False,
        ),
        ToolParameter(
            name="document_title",
            type="string",
            description="文档标题。默认从模板读取。",
            required=False,
        ),
    ]
    thinking_hint = "正在生成 Word 文档…"

    async def execute(
        self,
        tool_call,
        context: dict[str, Any] | None = None,
    ) -> AgentToolResult:
        args = tool_call.arguments or {}
        project_path_str = args.get("project_path", "").strip()
        filled_key = args.get("filled_template_artifact_key", "filled_template").strip()
        output_key = args.get("output_artifact_key", "draft_word").strip()
        doc_title = args.get("document_title", "").strip()

        if not project_path_str:
            return AgentToolResult.error_result(tool_call.id, "project_path is required")

        project_path = Path(project_path_str)
        store = ProjectStore()
        manifest = store.get_project(project_path)
        if manifest is None:
            return AgentToolResult.error_result(tool_call.id, f"Project not found: {project_path_str}")

        # 读取填充模板
        filled_meta = manifest.get_artifact(filled_key)
        if filled_meta is None:
            return AgentToolResult.error_result(
                tool_call.id, f"Filled template artifact '{filled_key}' not found."
            )

        filled_content = read_artifact(project_path, filled_meta, as_text=True)
        if filled_content is None:
            return AgentToolResult.error_result(
                tool_call.id, f"Filled template file not found: {filled_meta.path}"
            )

        try:
            filled_data = json.loads(filled_content)
        except json.JSONDecodeError as e:
            return AgentToolResult.error_result(
                tool_call.id, f"Invalid filled template JSON: {e}"
            )

        title = doc_title or filled_data.get("document_title", "投标材料")

        # 将模板数据转为 Markdown（作为 docx 生成的中间格式）
        md_content = _template_to_markdown(filled_data, title)

        # 保存 Markdown 中间产物
        md_path = "drafts/tender_draft.md"
        md_meta = ArtifactMeta(
            artifact_id=f"{output_key}_md",
            kind="word_draft_md",
            file_format="md",
            path=md_path,
            created_by=self.name,
            source_artifacts=[filled_key],
        )
        write_artifact(project_path, md_meta, content=md_content)
        store.register_artifact(project_path, f"{output_key}_md", md_meta)

        # 生成 docx：写入 output，作为用户可预览/下载的交付物
        docx_path = "output/tender_draft.docx"
        docx_abs_path = project_path / docx_path

        try:
            _generate_docx_from_markdown(
                md_content=md_content,
                output_path=docx_abs_path,
                title=title,
                project_path=project_path,
            )
        except Exception as e:
            logger.warning("[DocgenWordGenerate] docx generation failed: %s", e)
            docx_abs_path = None

        if docx_abs_path is None or not docx_abs_path.exists() or docx_abs_path.stat().st_size == 0:
            return AgentToolResult.error_result(
                tool_call.id,
                "Word generation failed: no .docx file was created.",
            )

        # 注册 artifact
        meta = ArtifactMeta(
            artifact_id=output_key,
            kind="word_draft",
            file_format="docx",
            path=docx_path,
            created_by=self.name,
            source_artifacts=[filled_key],
            metadata={"title": title},
        )

        store.register_artifact(project_path, output_key, meta)
        status_msg = "Word 文档生成成功"
        artifact_abs = str(docx_abs_path)

        append_event(
            project_path,
            "word_generated",
            step=self.name,
            artifact_id=output_key,
            format=meta.file_format,
        )

        logger.info("[DocgenWordGenerate] %s: %s", status_msg, artifact_abs)

        content = (
            f"{status_msg}:\n"
            f"  标题: {title}\n"
            f"  格式: {meta.file_format}\n"
            f"  路径: {meta.path}\n"
            f"  章节数: {len(filled_data.get('sections', []))}"
        )

        result = AgentToolResult(
            tool_call_id=tool_call.id,
            result_type=ToolResultType.TEXT,
            content=content,
            is_error=False,
            metadata={
                "artifact_id": output_key,
                "artifact_path": artifact_abs,
                "state_delta": {
                    f"docgen_state.artifacts.{output_key}": artifact_abs,
                    "docgen_state.draft_word_artifact_id": output_key,
                },
            },
            events=[],
        )

        digest = FrontendDigest(
            tool_name="docgen_word_standard_generate",
            display_type=ToolDisplayType.FILE,
            minimal=MinimalView(
                title=title,
                summary=f"Word 文档已生成 ({meta.file_format})",
                icon="wordmaster",
                status="success",
            ),
            detailed=DetailedView(
                title="Word 文档生成",
                sections=[
                    ViewSection(heading="摘要", content_type="text", data=content),
                ],
            ),
        )
        apply_dual_layer(result, digest, content)

        return result


def _template_to_markdown(filled_data: dict[str, Any], title: str) -> str:
    """将填充模板转为 Markdown"""
    lines = [f"# {title}", ""]

    for sec in filled_data.get("sections", []):
        level = sec.get("level", 1)
        sec_title = sec.get("title", "")
        content = sec.get("content", "") or sec.get("rag_result", "") or ""

        prefix = "#" * min(level + 1, 6)  # H2 ~ H6
        lines.append(f"{prefix} {sec_title}")
        lines.append("")
        if content:
            lines.append(content)
            lines.append("")

    return "\n".join(lines)


def _generate_docx_from_markdown(
    md_content: str,
    output_path: Path,
    title: str,
    project_path: Path,
) -> None:
    """使用 wordmaster 脚本生成 docx

    尝试调用 wordmaster 的 docx-js 脚本生成 Word 文档。
    """
    import os
    import subprocess
    import sys

    use_soffice = os.getenv("DOCGEN_USE_SOFFICE", "").lower() in {"1", "true", "yes"}
    if not use_soffice:
        _generate_docx_python(md_content, output_path, title)
        return

    js_script = _build_docx_js(md_content, title)
    script_path = project_path / "scripts" / "generate_draft.js"
    script_path.parent.mkdir(parents=True, exist_ok=True)
    script_path.write_text(js_script, encoding="utf-8")
    from egis_agent_plugins.core.tools.wordmaster._base import SCRIPTS_DIR

    try:
        result = subprocess.run(
            [sys.executable, str(SCRIPTS_DIR / "office" / "soffice.py"),
             "--input", str(script_path),
             "--output", str(output_path)],
            capture_output=True,
            text=True,
            timeout=120,
        )
        if result.returncode != 0:
            raise RuntimeError(f"docx generation script failed: {result.stderr[:500]}")
        if not output_path.exists() or output_path.stat().st_size == 0:
            raise RuntimeError("docx generation script completed but produced no output")
    except Exception as e:
        logger.warning("[DocgenWordGenerate] primary docx path failed, fallback to local docx: %s", e)
        _generate_docx_python(md_content, output_path, title)


def _generate_docx_python(md_content: str, output_path: Path, title: str) -> None:
    """使用 python-docx 简单生成 Word 文档"""
    try:
        from docx import Document

        doc = Document()
        doc.add_heading(title, level=0)

        for line in md_content.split("\n"):
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                level = sum(1 for c in stripped if c == "#")
                heading_text = stripped.lstrip("#").strip()
                if level == 1 and heading_text == title:
                    continue
                doc.add_heading(heading_text, level=min(level, 4))
            else:
                doc.add_paragraph(stripped)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        logger.info("[DocgenWordGenerate] python-docx generated: %s", output_path)
    except ImportError:
        logger.warning("[DocgenWordGenerate] python-docx not installed, using stdlib OOXML fallback")
        _generate_docx_minimal(md_content, output_path, title)


def _generate_docx_minimal(md_content: str, output_path: Path, title: str) -> None:
    """Generate a valid .docx using only Python stdlib."""
    import zipfile
    from xml.sax.saxutils import escape

    def paragraph(text: str, style: str | None = None) -> str:
        style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
        safe_text = escape(text)
        return (
            "<w:p>"
            f"{style_xml}"
            "<w:r>"
            '<w:rPr><w:rFonts w:eastAsia="SimSun" w:ascii="Arial" w:hAnsi="Arial"/></w:rPr>'
            f"<w:t>{safe_text}</w:t>"
            "</w:r>"
            "</w:p>"
        )

    body_parts = [paragraph(title, "Title")]
    for line in md_content.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("#"):
            level = min(stripped.count("#"), 4)
            heading_text = stripped.lstrip("#").strip()
            if heading_text and heading_text != title:
                body_parts.append(paragraph(heading_text, f"Heading{level}"))
        else:
            body_parts.append(paragraph(stripped))

    document_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:body>
    %s
    <w:sectPr>
      <w:pgSz w:w="11906" w:h="16838"/>
      <w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440" w:header="720" w:footer="720" w:gutter="0"/>
    </w:sectPr>
  </w:body>
</w:document>
""" % "\n".join(body_parts)

    styles_xml = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:style w:type="paragraph" w:default="1" w:styleId="Normal">
    <w:name w:val="Normal"/>
    <w:rPr><w:rFonts w:eastAsia="SimSun" w:ascii="Arial" w:hAnsi="Arial"/><w:sz w:val="24"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Title">
    <w:name w:val="Title"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr><w:spacing w:after="240"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="36"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading1">
    <w:name w:val="heading 1"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr><w:outlineLvl w:val="0"/><w:spacing w:before="240" w:after="120"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="32"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading2">
    <w:name w:val="heading 2"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr><w:outlineLvl w:val="1"/><w:spacing w:before="200" w:after="100"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="28"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading3">
    <w:name w:val="heading 3"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr><w:outlineLvl w:val="2"/><w:spacing w:before="160" w:after="80"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="26"/></w:rPr>
  </w:style>
  <w:style w:type="paragraph" w:styleId="Heading4">
    <w:name w:val="heading 4"/>
    <w:basedOn w:val="Normal"/>
    <w:pPr><w:outlineLvl w:val="3"/><w:spacing w:before="120" w:after="60"/></w:pPr>
    <w:rPr><w:b/><w:sz w:val="24"/></w:rPr>
  </w:style>
</w:styles>
"""

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(output_path, "w", compression=zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
  <Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
  <Default Extension="xml" ContentType="application/xml"/>
  <Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
  <Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/>
</Types>
""")
        zf.writestr("_rels/.rels", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">
  <Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/>
</Relationships>
""")
        zf.writestr("word/_rels/document.xml.rels", """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>
""")
        zf.writestr("word/document.xml", document_xml)
        zf.writestr("word/styles.xml", styles_xml)

    logger.info("[DocgenWordGenerate] stdlib docx generated: %s", output_path)


def _build_docx_js(md_content: str, title: str) -> str:
    """生成 docx-js 的 Node.js 脚本（占位，后续完善）"""
    return f"""
// Auto-generated docx-js script
// Title: {title}
const md_content = {json.dumps(md_content, ensure_ascii=False)};
console.log("docx-js generation placeholder");
"""
